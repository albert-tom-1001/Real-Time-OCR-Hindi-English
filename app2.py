import os
import tempfile
import streamlit as st
import pdf2image
import pytesseract
from PIL import Image
from docx import Document
import io
import cv2
import numpy as np
from googletrans import Translator

# Page configuration
st.set_page_config(
    page_title="Hindi to English OCR Platform",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton > button {
        width: 100%;
        margin-top: 1rem;
    }
    .upload-text {
        text-align: center;
        padding: 2rem;
    }
    </style>
""", unsafe_allow_html=True)

class OCRPlatform:
    def __init__(self, output_format='docx'):
        self.output_format = output_format
        self.translator = Translator()

    def extract_text_from_pdf(self, pdf_path, page_range=None):
        try:
            # Convert PDF to images
            images = pdf2image.convert_from_path(
                pdf_path,
                dpi=300,
                thread_count=2,
                first_page=page_range[0] if page_range else None,
                last_page=page_range[1] if page_range else None
            )
            
            if not images:
                st.error("No images were extracted from the PDF. Please check if the file is valid.")
                return None

            # Process images and extract text
            with st.spinner('Processing images...'):
                processed_images = self.preprocess_images(images)
                extracted_text = self.perform_ocr(processed_images)
                
            # Translate text
            with st.spinner('Translating text...'):
                translated_text = [self.translate_text(text) for text in extracted_text if text.strip()]

            if self.output_format == 'docx':
                return self.generate_docx(translated_text)
            else:
                return '\n\n'.join(translated_text)

        except Exception as e:
            st.error(f"An error occurred during processing: {str(e)}")
            return None

    def preprocess_images(self, images):
        processed_images = []
        progress_bar = st.progress(0)
        
        for idx, image in enumerate(images):
            # Convert to grayscale
            image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
            
            # Apply adaptive thresholding
            image_cv = cv2.adaptiveThreshold(
                image_cv, 255, 
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # Denoise
            image_cv = cv2.fastNlMeansDenoising(image_cv, h=10)
            
            # Convert back to PIL Image
            processed_images.append(Image.fromarray(image_cv))
            
            # Update progress
            progress_bar.progress((idx + 1) / len(images))
        
        progress_bar.empty()
        return processed_images

    def perform_ocr(self, images):
        extracted_text = []
        progress_bar = st.progress(0)
        
        for idx, image in enumerate(images):
            text = pytesseract.image_to_string(image, lang='hin')
            if text.strip():
                extracted_text.append(text.strip())
            progress_bar.progress((idx + 1) / len(images))
        
        progress_bar.empty()
        return extracted_text

    def translate_text(self, text):
        try:
            if not text.strip():
                return ""
            translated = self.translator.translate(text, src='hi', dest='en')
            return translated.text
        except Exception as e:
            st.warning(f"Translation error: {str(e)}")
            return text

    def generate_docx(self, translated_text):
        document = Document()
        
        # Add title
        document.add_heading('Translated Document', 0)
        
        # Add translation date
        from datetime import datetime
        document.add_paragraph(f'Translation Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add content
        for page_text in translated_text:
            if page_text.strip():
                document.add_paragraph(page_text)
                document.add_page_break()

        # Save to bytes
        doc_bytes = io.BytesIO()
        document.save(doc_bytes)
        return doc_bytes.getvalue()

def main():
    st.title("📝 Hindi to English OCR Platform")
    st.markdown("### Convert Hindi PDF documents to English text")
    
    # Sidebar configurations
    st.sidebar.title("⚙️ Settings")
    
    # Output format selection
    output_format = st.sidebar.selectbox(
        "Select output format",
        ["docx", "text"],
        help="Choose the format for the translated output"
    )
    
    # File uploader
    uploaded_file = st.file_uploader(
        "Upload your Hindi PDF document",
        type="pdf",
        help="Upload a PDF file containing Hindi text"
    )

    if uploaded_file is not None:
        try:
            # Create temporary directory
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save uploaded file
                temp_pdf_path = os.path.join(temp_dir, uploaded_file.name)
                with open(temp_pdf_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())

                # Get PDF info
                info = pdf2image.pdfinfo_from_path(temp_pdf_path)
                max_pages = info["Pages"]

                # Page range selection
                st.sidebar.markdown("### Page Selection")
                page_range = st.sidebar.text_input(
                    f"Enter page range (1-{max_pages})",
                    value="1",
                    help=f"Enter a single page number or range (e.g., 1-3). Max: {max_pages}"
                )

                try:
                    if '-' in page_range:
                        start_page, end_page = map(int, page_range.split("-"))
                        if end_page > max_pages:
                            st.sidebar.error(f"Maximum page number is {max_pages}")
                            return
                    else:
                        start_page = end_page = int(page_range)
                        if start_page > max_pages:
                            st.sidebar.error(f"Maximum page number is {max_pages}")
                            return
                except ValueError:
                    st.sidebar.error("Please enter valid page numbers")
                    return

                # Process button
                if st.button("Process Document"):
                    with st.spinner('Processing your document... Please wait.'):
                        ocr_platform = OCRPlatform(output_format=output_format)
                        extracted_content = ocr_platform.extract_text_from_pdf(
                            temp_pdf_path,
                            (start_page, end_page)
                        )

                        if extracted_content:
                            st.success("Processing completed!")
                            
                            if output_format == 'docx':
                                st.download_button(
                                    label="📥 Download DOCX",
                                    data=extracted_content,
                                    file_name=f"{uploaded_file.name.split('.')[0]}_translated.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                )
                            else:
                                st.markdown("### Translated Text:")
                                st.text_area("", value=extracted_content, height=300)
                                
                                # Copy button
                                if st.button("📋 Copy to Clipboard"):
                                    st.write("Text copied to clipboard!")
                                    st.session_state['clipboard'] = extracted_content

        except Exception as e:
            st.error("An error occurred while processing the document.")
            st.error(f"Error details: {str(e)}")
            st.info("Please try again with a different PDF or contact support if the issue persists.")

    # Instructions
    else:
        st.markdown("""
        ### Instructions:
        1. Upload a PDF document containing Hindi text
        2. Select the output format (DOCX or Text)
        3. Specify the page range to process
        4. Click 'Process Document' to start conversion
        
        ### Notes:
        - The processing time depends on the document size
        - For best results, use clear, well-scanned documents
        - Large documents may take several minutes to process
        """)

if __name__ == "__main__":
    main()
